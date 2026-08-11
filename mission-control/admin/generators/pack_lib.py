#!/usr/bin/env python3
"""
Investor pack DOCX library (PDS + SA) for Mission Control.

Absorbed from retired `_tools/investor-pack-wizard/`.
Call via Mission Control HLT Detail → Generate draft pack.
Standalone wizard UI is gone — do not re-add as a separate tool.

DOCX export for counsel review →
  _assets/horses/{evo_slug}/documents/investor-packs/
"""

from __future__ import annotations

import argparse
import cgi
import json
import mimetypes
import os
import re
import socketserver
import sys
import tempfile
from datetime import date, datetime, timezone
from http.server import SimpleHTTPRequestHandler
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlparse

# generators → admin → mission-control → _tools → evo_01
HERE = Path(__file__).resolve().parent
ROOT = HERE.parents[3]
ASSETS_HORSES = ROOT / "_assets" / "horses"
WEB_PUBLIC = ROOT / "02_website" / "public"

EVO_MAP = {
    "tml-x-yearn": "turn-me-loose-x-yearn",
    "nellie": "almanzor-x-night-danza",
    "i-stole-a-manolo": "i-stole-a-manolo",
    "first-gear": "first-gear",
    "hottathanafantasy": "hottathanafantasy",
    "prudentia": "prudentia",
}

CATEGORIES = ("vet", "valuation", "risk", "whitepaper", "other", "investor-packs")
SAFE_NAME = re.compile(r"[^A-Za-z0-9._\- ()\[\]]+")


def evo_slug(web_slug: str) -> str:
    return EVO_MAP.get(web_slug, web_slug)


def docs_root(web_slug: str) -> Path:
    return ASSETS_HORSES / evo_slug(web_slug) / "documents"


def safe_filename(name: str) -> str:
    base = Path(name).name
    base = SAFE_NAME.sub("_", base).strip("._ ")
    return base or "upload.bin"


def list_attachments(web_slug: str) -> list[dict]:
    root = docs_root(web_slug)
    out: list[dict] = []
    if not root.is_dir():
        return out
    for cat in CATEGORIES:
        d = root / cat
        if not d.is_dir():
            continue
        for f in sorted(d.rglob("*")):
            if not f.is_file():
                continue
            if f.suffix.lower() not in {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"}:
                continue
            rel = f.relative_to(ASSETS_HORSES)
            out.append(
                {
                    "category": cat,
                    "name": f.name,
                    "path": str(f.relative_to(ROOT)),
                    "url": "/horse-assets/" + "/".join(quote_seg(p) for p in rel.parts),
                    "label": f"[{cat}] {f.name}",
                }
            )
    # also loose files in documents/ root
    for f in sorted(root.iterdir()) if root.is_dir() else []:
        if f.is_file() and f.suffix.lower() in {".pdf", ".docx"}:
            rel = f.relative_to(ASSETS_HORSES)
            out.append(
                {
                    "category": "other",
                    "name": f.name,
                    "path": str(f.relative_to(ROOT)),
                    "url": "/horse-assets/" + "/".join(quote_seg(p) for p in rel.parts),
                    "label": f"[documents] {f.name}",
                }
            )
    return out


def quote_seg(p: str) -> str:
    from urllib.parse import quote

    return quote(p)


def _add_months(start: str, months) -> str:
    if not start or not months:
        return ""
    try:
        y, m, d = [int(x) for x in str(start).split("-")[:3]]
        months = int(months)
    except Exception:
        return ""
    m0 = m - 1 + months
    y += m0 // 12
    m = m0 % 12 + 1
    if m == 1:
        ey, em = y - 1, 12
    else:
        ey, em = y, m - 1
    return f"{ey:04d}-{em:02d}-{min(d, 28):02d}"


def _fmt_date(iso: str) -> str:
    if not iso:
        return "—"
    try:
        y, m, d = [int(x) for x in str(iso).split("-")[:3]]
        months = [
            "",
            "January",
            "February",
            "March",
            "April",
            "May",
            "June",
            "July",
            "August",
            "September",
            "October",
            "November",
            "December",
        ]
        return f"{d} {months[m]} {y}"
    except Exception:
        return str(iso)


def _money(n) -> str:
    try:
        x = float(n)
        if x == int(x):
            return f"{int(x):,}"
        return f"{x:,.2f}"
    except Exception:
        return str(n or "—")


def _resolve_cover_path(url: str) -> Path | None:
    """Map wizard URL (/horse-assets/… or /web-public/…) to a local file."""
    if not url:
        return None
    u = url.split("?", 1)[0]
    if u.startswith("/horse-assets/"):
        p = ASSETS_HORSES / unquote(u[len("/horse-assets/") :])
    elif u.startswith("/web-public/"):
        p = WEB_PUBLIC / unquote(u[len("/web-public/") :])
    elif u.startswith("_assets/") or u.startswith("02_website/"):
        p = ROOT / u
    else:
        return None
    return p if p.is_file() else None


def _set_narrow_margins(doc, inches: float = 0.85):
    from docx.shared import Inches

    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _style_doc(doc):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for level in range(1, 4):
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Times New Roman"
            h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if level == 1:
                h.font.size = Pt(16)
            elif level == 2:
                h.font.size = Pt(13)
            else:
                h.font.size = Pt(11)
        except Exception:
            pass


def _p(doc, text: str, *, bold: bool = False, italic: bool = False, size: int | None = None, space_after: int = 8):
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def _bullets(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _footer_line(doc, left: str, right: str):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(f"{left}  ·  {right}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(payload: dict) -> Path:
    """
    Write investor-facing pack DOCX (same shape as final investor PDF):
    Cover → Part A PDS → Part B SA → Appendices.
    Marked DRAFT until counsel freezes version.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    horse = payload.get("horse") or {}
    inv = payload.get("investor") or {}
    atts = payload.get("attachments") or []
    web_slug = horse.get("slug") or "unknown"
    evo = evo_slug(web_slug)
    out_dir = ASSETS_HORSES / evo / "documents" / "investor-packs"
    out_dir.mkdir(parents=True, exist_ok=True)

    name = horse.get("display_name") or horse.get("legal_name") or web_slug
    legal = horse.get("legal_name") or name
    inv_bit = ""
    if inv.get("name"):
        inv_bit = "-" + re.sub(r"\s+", "_", inv["name"])[:40]
    ts = datetime.now(timezone.utc).strftime("%Y%m%d")
    out_path = out_dir / f"Investor-Pack-{web_slug}{inv_bit}-DRAFT-{ts}.docx"

    stake = float(horse.get("leasehold_stake_pct") or 0)
    shares = int(horse.get("shares_total") or 0) if horse.get("shares_total") not in ("", None) else 0
    lot = float(horse.get("lot_pct") or 0) or (stake / shares if shares else 0)
    months = horse.get("lease_period_months") or ""
    start = horse.get("lease_start_date") or ""
    end = horse.get("lease_end_date") or _add_months(start, months)
    price = horse.get("price_per_share_nzd")
    ret = float(horse.get("investor_return_pct") or 75)
    owner_keep = 100 - ret
    fee = horse.get("platform_fee_pct") or 5
    owner_rate = horse.get("owner_rate_per_1pct_month")
    about = horse.get("about_horse") or horse.get("story") or ""
    race = horse.get("race_schedule") or "Updates via Evolution investor channels as the horse progresses."
    doc_date = payload.get("document_date") or date.today().isoformat()
    inv_name = inv.get("name") or ""
    inv_pct = inv.get("interest_pct")
    inv_tokens = inv.get("tokens")
    if inv_pct not in (None, "") and (inv_tokens in (None, "")) and lot:
        try:
            inv_tokens = round(float(inv_pct) / lot, 2)
        except Exception:
            pass
    if inv_pct not in (None, ""):
        inv_label = f"{inv_pct}%"
        if inv_tokens not in (None, ""):
            inv_label += f" ({inv_tokens} token{'s' if float(inv_tokens) != 1 else ''})"
    else:
        inv_label = "— (allocation not set — general syndicate pack)"

    nick = name != legal

    doc = Document()
    _set_narrow_margins(doc, 0.9)
    _style_doc(doc)

    # ─── Cover ───
    banner = doc.add_paragraph()
    br = banner.add_run(
        "DRAFT — Investor Pack for review. Not an executed legal document until counsel sign-off and final PDF freeze."
    )
    br.bold = True
    br.font.size = Pt(9)
    br.font.color.rgb = RGBColor(0x5C, 0x4A, 0x00)
    banner.paragraph_format.space_after = Pt(14)

    meta = doc.add_paragraph()
    mr = meta.add_run("Evolution Stables  ·  NZTR Authorised Syndicator")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    h = doc.add_heading("Investor Pack", level=0)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph()
    r = sub.add_run(f"{name} Syndicate")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph("Product Disclosure Statement + Syndicate Agreement")
    doc.add_paragraph(f"{_money(stake).rstrip('0').rstrip('.') if isinstance(stake, float) else stake}% Leasehold Stake in {legal} (NZ)")

    cover_path = _resolve_cover_path(horse.get("cover_image_url") or "")
    if cover_path:
        try:
            doc.add_picture(str(cover_path), width=Inches(5.8))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            _p(doc, f"[Cover image could not be embedded: {cover_path.name}]", italic=True, size=9)

    _p(doc, f"Date: {_fmt_date(doc_date)}", size=10)
    _p(doc, "Evolution Stables\nalex@evolutionstables.nz | +64 21 0828 0901\nevolutionstables.nz", size=10)
    if inv_name:
        _p(doc, f"Prepared for: {inv_name}\nInterest: {inv_label}", bold=True, size=11)

    box = doc.add_paragraph()
    box.add_run("This pack includes:\n").bold = True
    box.add_run(
        "• Part A — Product Disclosure Statement (PDS) §§1–21\n"
        "• Part B — Syndicate Agreement (SA)\n"
        "• Appendices — supporting documents (vet, valuation, other)"
    )
    _footer_line(doc, "evolutionstables.nz", f"{name} Syndicate · DRAFT")
    doc.add_page_break()

    # ─── PART A PDS ───
    _p(doc, "PART A", bold=True, size=10, space_after=4)
    doc.add_heading("Product Disclosure Statement", level=1)
    _p(doc, f"{name} Syndicate", bold=True)
    _p(doc, f"{stake:g}% Leasehold Stake in {legal} (NZ)")
    if nick:
        _p(doc, f"Public name / nickname: {name}  |  Pedigree / legal description: {legal}", size=10)
    _p(doc, f"Date: {_fmt_date(doc_date)}", size=10)

    doc.add_heading("A New Way to Own", level=2)
    doc.add_paragraph(
        "Evolution Stables is a New Zealand–based, NZTR-authorised syndicator offering fractional, "
        "fixed-term digital leasehold interests in thoroughbreds. We partner with leading owners and trainers "
        "to lease premium bloodstock, then make those opportunities available in fractional digital shares."
    )
    doc.add_paragraph(
        "To deliver this securely and at scale, we work with Tokinvest, a licensed and regulated platform "
        "(Dubai VARA-approved) that manages investor onboarding, compliance, and secure digital issuance "
        "of syndicate tokens."
    )
    doc.add_paragraph(
        "Our approach blends the tradition, passion, and community of racing with modern technology "
        "and transparent processes—helping to build the next generation of owners."
    )
    _p(doc, "Alex Baddeley\nFounder, Evolution Stables", bold=True, size=10)

    doc.add_heading("Key Information Summary", level=2)
    doc.add_heading("What is this?", level=3)
    doc.add_paragraph(
        f"This Product Disclosure Statement outlines a {months}-month leasehold interest in {legal}"
        + (f" (marketed as {name})" if nick else "")
        + f", a New Zealand thoroughbred. Evolution Stables has secured a {stake:g}% leasehold stake "
        "and is making fractional participation available through its digital-syndication model."
    )
    doc.add_heading("Who is this for?", level=3)
    doc.add_paragraph(
        "Individuals seeking exposure to racehorse ownership in a simplified, digital format—no prior experience required."
    )
    doc.add_heading("How does it work?", level=3)
    doc.add_paragraph(
        f"You lease a share of our {stake:g}% stake in {name}. Your one-time fee covers costs for the term as disclosed. "
        f"You receive {ret:g}% of revenue generated from racing, sponsorship, or commercial activity "
        "attributable to the syndicate stake, proportional to your share."
    )
    doc.add_heading("Structure", level=3)
    _bullets(
        doc,
        [
            f"Split into {shares} digital tokens (each representing {lot:g}% of the horse)",
            f"Fixed for the term {_fmt_date(start)} to {_fmt_date(end)} ({months} months)",
            "Fully managed by Evolution Stables",
            "Digitally issued and administered via Tokinvest (regulated by VARA, Dubai) where applicable",
        ],
    )
    doc.add_heading("Fees", level=3)
    _bullets(
        doc,
        [
            f"${_money(price)} NZD per {lot:g}% token (lot)",
            f"Platform fee on list rate: {fee}%",
            f"Owner rate: ${_money(owner_rate)} / month / 1% (if applicable)",
        ],
    )
    doc.add_heading("Returns", level=3)
    doc.add_paragraph(
        "Returns are variable, based on race performance and commercial revenue. There are no guarantees, "
        "and you may not recover your original investment."
    )
    doc.add_heading("Race schedule expectation", level=3)
    doc.add_paragraph(str(race))
    doc.add_heading("Risks", level=3)
    _bullets(
        doc,
        [
            "The horse may underperform, suffer injury, or be retired early.",
            "Early termination may result in a pro-rata refund where applicable.",
            "Limited liquidity—resale may not be immediate or available.",
            "You could lose the full amount of your investment.",
        ],
    )
    doc.add_heading("How to invest", level=3)
    doc.add_paragraph(
        "Applications via Evolution Stables / the marketplace. Identity verification may be required under AML/CFT laws. "
        "Purchases only when the campaign is listed and the site allows charging."
    )
    doc.add_heading("Manager", level=3)
    doc.add_paragraph(
        "Evolution Stables Ltd (NZBN: 9429050177875) is an authorised NZTR syndicator specialising in "
        "digital-syndication of leasehold interests in New Zealand thoroughbreds."
    )
    _footer_line(doc, "evolutionstables.nz", f"{name} Syndicate · PDS")
    doc.add_page_break()

    doc.add_heading("Contents", level=2)
    for i, title in enumerate(
        [
            "About Evolution Stables",
            f"About the Horse – {name}",
            "Upfront Costs and Expenses",
            "Revenue Streams",
            "Ongoing Costs and Expenses",
            "Minimum Amount",
            "What Taxes Will You Pay?",
            "Insurance Information",
            "Valuation Reports",
            "Veterinary Report",
            "Material Interests and Commissions",
            "Risk Disclosure",
            "Responsible Investment",
            "Legal Structure and Oversight",
            "Termination & Forfeiture",
            "Transfer of Interests",
            "Records and Administration",
            "How to Complain",
            "Investment Details",
            "Investor Declaration",
            "Promoter Declaration",
        ],
        1,
    ):
        doc.add_paragraph(f"{i}. {title}")

    doc.add_heading("1. About Evolution Stables", level=2)
    doc.add_paragraph(
        "Evolution Stables is a licensed NZTR-authorised syndicator reshaping how people experience racehorse ownership. "
        "Fractional ownership, simple onboarding, and clear timeframes—enabled by digital infrastructure. "
        "We work closely with NZTR, Tokinvest (VARA), and leading trainers and bloodstock partners."
    )

    doc.add_heading(f"2. About the Horse – {name}", level=2)
    if about:
        for para in str(about).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        _p(doc, "Narrative to be completed.", italic=True)
    doc.add_heading("Key Details", level=3)
    details = [
        f"Display name: {name}",
        f"Legal / pedigree name: {legal}",
        f"Foaling date: {_fmt_date(horse.get('foaling_date') or '')}",
        f"Colour: {horse.get('colour') or '—'}",
        f"Sex: {horse.get('sex') or '—'}",
        f"Height: {horse.get('height') or '—'}",
        f"Sire: {horse.get('sire_name') or '—'}",
        f"Dam: {horse.get('dam_name') or '—'}",
        f"Breeder: {horse.get('breeder') or '—'}",
        f"Trainer: {horse.get('trainer_name') or '—'}",
        f"Training base: {horse.get('trainer_stable') or '—'} — {horse.get('trainer_location') or '—'}",
        f"Life number: {horse.get('life_number') or 'Pending / not registered'}",
        f"Microchip: {horse.get('microchip') or 'Pending registration'}",
        f"loveracing.nz ID: {horse.get('loveracing_id') or '—'}",
        f"Identity status: {horse.get('identity_status') or '—'}",
        f"Lease duration: {months} months ({_fmt_date(start)} – {_fmt_date(end)})",
        f"Stake: {stake:g}% leasehold interest held by Evolution Stables",
        f"Listing: {shares} digital tokens ({lot:g}% each)",
        f"Lessor / owner: {horse.get('owner_name') or '—'}",
        f"Record: {horse.get('record') or ('Wins ' + str(horse.get('wins') or 0) + ', placed ' + str(horse.get('placed') or 0))}",
    ]
    _bullets(doc, details)

    doc.add_heading("3. Upfront Costs and Expenses", level=2)
    doc.add_paragraph(
        f"The {name} Syndicate is structured as a fixed-cost lease for the disclosed term. "
        "You pay one upfront amount to participate (unless renewal is agreed later)."
    )
    _p(doc, f"Total lease cost: ${_money(price)} NZD per {lot:g}% token", bold=True)
    doc.add_paragraph(
        "This cost is intended to cover training and care, syndicate management and reporting, "
        "NZTR registration/compliance, Tokinvest platform administration, and routine costs associated "
        "with the leasehold interest for the term — as finalised in the executed PDS."
    )

    doc.add_heading("4. Revenue Streams", level=2)
    doc.add_paragraph(
        f"Under the lease arrangement, {ret:g}% of gross revenue attributable to the syndicate stake "
        f"is allocated to token holders, with the remaining {owner_keep:g}% retained by the lessor/structure as disclosed. "
        "Revenue may include prizemoney, sponsorship, media/commercial rights, and other verifiable income "
        "attributable to the horse during the lease period."
    )

    doc.add_heading("5. Ongoing Costs and Expenses", level=2)
    doc.add_paragraph(
        "Unless otherwise stated in the final PDS, routine costs for the disclosed term are included in the upfront fee. "
        "Renewals or extraordinary costs (if any) will be disclosed before they apply."
    )

    doc.add_heading("6. Minimum Amount", level=2)
    doc.add_paragraph(
        f"The minimum investment is one digital token ({lot:g}% of the horse), priced at ${_money(price)} NZD, subject to availability."
    )

    doc.add_heading("7. What Taxes Will You Pay?", level=2)
    doc.add_paragraph(
        f"Revenue earned from your participation in the {name} Syndicate may be subject to tax, "
        "depending on your personal circumstances and jurisdiction. Evolution Stables does not provide tax advice. "
        "You should consult a qualified adviser."
    )

    doc.add_heading("8. Insurance Information", level=2)
    doc.add_paragraph(
        "Insurance arrangements for the horse (mortality, loss of use, or other cover), if any, are held "
        "or arranged by the owner/lessor or as disclosed in the final pack schedules. Leaseholders should "
        "not assume they are named insureds unless expressly stated."
    )

    doc.add_heading("9. Valuation Reports", level=2)
    doc.add_paragraph(
        "Any formal valuation of the horse or the leasehold interest will be appended as a schedule when available. "
        "Until then, the offer price of tokens is a fixed commercial term of the lease syndicate and is not a "
        "representation of residual freehold value of the horse."
    )

    doc.add_heading("10. Veterinary Report", level=2)
    doc.add_paragraph(
        f"A current veterinary assessment suitable for syndication purposes should be attached as an appendix "
        f"in the executed pack. Until attached, treat fitness for training/racing as subject to that report. "
        f"{name} is described for marketing and education based on information supplied by the trainer/owner; "
        "it is not a veterinary certificate."
    )

    doc.add_heading("11. Material Interests and Commissions", level=2)
    doc.add_paragraph(
        "Evolution Stables Ltd acts as promoter and manager and receives management and platform-related fees "
        "as disclosed. The lessor "
        f"({horse.get('owner_name') or 'as disclosed'}) retains interests in the horse outside the syndicate stake."
    )

    doc.add_heading("12. Risk Disclosure", level=2)
    _bullets(
        doc,
        [
            "Racing is speculative; you may lose some or all of your capital.",
            "No guarantee of race starts, placings, wins, or commercial returns.",
            "Injury, illness, or early retirement may reduce or end racing activity.",
            "Liquidity is limited; resale of tokens may not be available or timely.",
            "Participation is leasehold only — not freehold ownership of the horse.",
            "Past performance of sire, dam, trainer, or related horses is not a reliable indicator of future results.",
        ],
    )

    doc.add_heading("13. Responsible Investment", level=2)
    doc.add_paragraph(
        "Only invest amounts you can afford to lose. Syndicate participation is intended as a recreational "
        "interest in thoroughbred racing with financial risk. Seek independent advice if unsure."
    )

    doc.add_heading("14. Legal Structure and Oversight", level=2)
    doc.add_paragraph(
        "The syndicate is promoted under NZTR Bloodstock Syndication Code of Practice frameworks. "
        "Evolution Stables Ltd (NZBN: 9429050177875) is the promoter/manager. Digital issuance/administration "
        "may be supported by Tokinvest under applicable VARA-related arrangements."
    )

    doc.add_heading("15. Termination & Forfeiture", level=2)
    doc.add_paragraph(
        "The lease ends on the stated end date, or earlier if the horse is retired, sold under disclosed "
        "arrangements, or as otherwise provided in this PDS, the Syndicate Agreement, and NZTR rules. "
        "Early termination may involve pro-rata treatment of fees where disclosed."
    )

    doc.add_heading("16. Transfer of Interests", level=2)
    doc.add_paragraph(
        "Tokens may only be transferred with Manager consent and in accordance with NZTR and platform rules. "
        "Transfers must be recorded by the Manager."
    )

    doc.add_heading("17. Records and Administration", level=2)
    doc.add_paragraph(
        "The Manager will maintain records of syndicate membership, communications, and distributions as required "
        "for NZTR and internal administration. Updates on the horse will be provided through Evolution channels "
        "while the lease is active."
    )

    doc.add_heading("18. How to Complain", level=2)
    doc.add_paragraph(
        "Complaints may be directed first to Evolution Stables (alex@evolutionstables.nz). If unresolved, "
        "participants may use applicable NZTR or other regulatory complaint pathways as advised in the final pack."
    )

    doc.add_heading("19. Investment Details", level=2)
    doc.add_paragraph(f"I hereby apply for the following interest in the {name} Syndicate:")
    _p(doc, inv_label, bold=True)
    if inv_name:
        doc.add_paragraph(f"Applicant: {inv_name}")
    else:
        _p(doc, "General pack — no named applicant yet.", italic=True)

    doc.add_heading("20. Investor Declaration", level=2)
    doc.add_paragraph("By signing this Application Form, I acknowledge and agree that:")
    for i, t in enumerate(
        [
            f"I have received, read, and understood the Product Disclosure Statement (PDS) and the Syndicate Agreement for the {name} Syndicate.",
            "I agree to be bound by the Syndicate Agreement, the PDS, and the Rules of Racing administered by NZTR.",
            "My participation is strictly as a leaseholder and confers no ownership rights in the horse.",
            "I understand the risks set out in the PDS, including that I may not recover my original investment.",
            "I consent to my details being provided to NZTR and Tokinvest for AML/CFT, registration, and regulatory purposes.",
        ],
        1,
    ):
        doc.add_paragraph(f"{i}. {t}")
    doc.add_paragraph(f"Name: {inv_name or '________________________'}")
    doc.add_paragraph("Signed: ____________________     Date: ____________________")

    doc.add_heading("21. Promoter Declaration", level=2)
    doc.add_paragraph(
        "I, Alex Baddeley, as the promoter of this syndicate, declare that the information provided in this "
        "Product Disclosure Statement is, to the best of my knowledge, true and correct, and that I am not aware "
        "of any information that would make this statement misleading in any material respect."
    )
    doc.add_paragraph("Signed: ____________________     Date: ____________________")
    doc.add_paragraph("Alex Baddeley\nDirector, Evolution Stables Ltd\n8 Huia Street, Auckland, New Zealand")
    _footer_line(doc, "evolutionstables.nz", f"{name} Syndicate · PDS · DRAFT")
    _p(doc, "End of Part A — Product Disclosure Statement. Syndicate Agreement follows.", italic=True, size=9)
    doc.add_page_break()

    # ─── PART B SA ───
    _p(doc, "PART B", bold=True, size=10, space_after=4)
    doc.add_heading("Syndicate Agreement", level=1)
    _p(doc, f"{name} Syndicate Agreement", bold=True)
    doc.add_paragraph(f"Evolution Stables – {legal} (NZ)")
    doc.add_paragraph("NZTR Authorised Syndicator")
    _p(doc, f"Date: {_fmt_date(doc_date)}", size=10)
    if inv_name:
        _p(doc, f"Prepared for: {inv_name} · {inv_label}", size=10)

    sa_sections = [
        (
            "1. Formation",
            "A Syndicate is formed under the New Zealand Thoroughbred Racing Inc. (\"NZTR\") Bloodstock "
            "Syndication Code of Practice (\"COP\"), in accordance with Clause 3.1, by the Promoter as set out "
            "in the attached Product Disclosure Statement (\"the Syndicate\").",
        ),
        (
            "2. Object",
            f"The object of the Syndicate is to lease and race {legal}"
            + (f" (marketed as {name})" if nick else "")
            + f" as a recreational pursuit, with all syndicate members holding digital leasehold shares for a fixed "
            f"{months}-month term, as described in the Product Disclosure Statement.",
        ),
        (
            "3. Agreement and Parties",
            "This Agreement is binding on the Promoter (Evolution Stables Ltd), the Syndicate Manager "
            "(Evolution Stables), and each Shareholder (token holder) as defined in the Product Disclosure Statement. "
            "By signing the Application Form (physically or digitally via Tokinvest), each Shareholder agrees to be "
            "bound by this Agreement and the Product Disclosure Statement. This Agreement may only be altered by "
            "special resolution (75% of shareholding) and must not increase a Shareholder's liability beyond what is "
            "disclosed in the Product Disclosure Statement.",
        ),
        (
            "4. Syndicate Shares (Tokens)",
            f"The Syndicate is divided into {shares} digital shares (tokens) of {lot:g}% each, representing a "
            f"{stake:g}% leasehold interest in {legal} for the Lease Term. Shares are issued and recorded on "
            "Tokinvest's VARA-compliant platform (where applicable), enabling digital onboarding, compliance, and "
            "(subject to Manager approval) transfer. Each Shareholder's rights and obligations are proportional to "
            "their shareholding.",
        ),
        (
            "5. Lease Duration",
            f"The lease term is fixed at {months} months, commencing {_fmt_date(start)} and ending {_fmt_date(end)}. "
            "All lease terms are counted in full calendar months. The lease may be renewed or extended at the "
            "Manager's and Lessor's discretion, with terms disclosed to Shareholders prior to renewal.",
        ),
        (
            "6. Manager's Powers and Duties",
            "The Manager (Evolution Stables) is responsible for overall lease administration and NZTR compliance, "
            "including communication with shareholders and coordination with licensed professionals. The Manager may: "
            "make day-to-day racing, training, spelling, and welfare decisions; appoint or change trainers in consultation "
            "with the owner or Racing Manager; deduct disclosed fees; delegate to licensed parties; and provide regular updates.",
        ),
        (
            "7. Financial Contributions and Fees",
            f"All costs for the disclosed term are covered by the one-time, upfront lease fee as detailed in the "
            f"Product Disclosure Statement (${_money(price)} NZD per {lot:g}% token), unless the lease is renewed. "
            "Tokinvest platform fees and management fees are included in the upfront payment as disclosed.",
        ),
        (
            "8. Revenue Streams and Distribution",
            f"Token holders are entitled to a share of revenue generated by {name} during the lease period "
            f"attributable to the syndicate stake, proportional to their leased stake. The disclosed revenue split is "
            f"{ret:g}% to token holders and {owner_keep:g}% retained by the lessor/structure as set out in the PDS. "
            "Potential streams include race winnings, sponsorship, media, digital, hospitality, and other verifiable income.",
        ),
        (
            "9. Insurance and Early Termination",
            "Insurance and early-termination consequences are as disclosed in the PDS. Where the lease ends early "
            "due to circumstances set out in the PDS, fee treatment follows the PDS.",
        ),
        (
            "10. Transfer of Shares (Tokens)",
            "Shares (tokens) may only be transferred with the written consent of the Manager and in accordance "
            "with NZTR and platform requirements. All transfers must be recorded by the Syndicate Manager.",
        ),
        (
            "11. Decision-Making and Voting",
            "Day-to-day racing and welfare decisions rest with the Manager (and licensed professionals as delegated). "
            "Matters requiring special resolution (75% of shareholding) are limited as described in clause 3.",
        ),
        (
            "12. Dispute Resolution",
            "Parties will first attempt good-faith resolution with the Manager. Unresolved disputes may proceed under "
            "New Zealand law and any dispute process referenced in the final executed pack or NZTR frameworks.",
        ),
        (
            "13. Removal of Manager",
            "Removal or replacement of the Manager, if ever applicable, will follow the process set out in the final "
            "executed Syndicate Agreement and NZTR requirements. Until then, Evolution Stables Ltd remains Manager.",
        ),
        (
            "14. Winding Up and Post-Lease Arrangements",
            "On expiry or winding up, remaining distributions (if any) and administrative close-out follow the PDS. "
            "Leasehold rights end on termination; no freehold interest in the horse is created by this Agreement.",
        ),
        (
            "15. Governing Law",
            "This Syndicate Agreement and the terms of the Product Disclosure Statement are governed by the "
            "laws of New Zealand and the Rules of Racing administered by NZTR.",
        ),
        (
            "16. Notices",
            "Notices to participants may be sent by email to the address provided on application, or via the "
            "Evolution / Tokinvest platforms used for the syndicate.",
        ),
        (
            "17. Execution",
            "This Agreement may be executed in counterparts and by electronic means. Together with the PDS and "
            "completed application, it forms the basis of the participant's leasehold interest.",
        ),
        (
            "18. Acknowledgement",
            "Each Shareholder acknowledges they have received the PDS and this Agreement, understand that "
            "participation is leasehold only (not freehold ownership of the horse), and accept the risks of racing "
            "investment including possible loss of capital.",
        ),
    ]
    for title, body in sa_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_paragraph("")
    _p(doc, "Shareholder", bold=True)
    doc.add_paragraph(f"Name: {inv_name or '________________________'}")
    doc.add_paragraph(f"Interest: {inv_label}")
    doc.add_paragraph("Signed: ____________________     Date: ____________________")
    doc.add_paragraph("")
    _p(doc, "For the Promoter / Manager — Evolution Stables Ltd", bold=True)
    doc.add_paragraph("Alex Baddeley, Director")
    doc.add_paragraph("Signed: ____________________     Date: ____________________")
    _footer_line(doc, "evolutionstables.nz", f"{name} Syndicate · SA · DRAFT")
    doc.add_page_break()

    # ─── Appendices ───
    doc.add_heading("Appendices — supporting documents", level=1)
    doc.add_paragraph(
        "The following documents are incorporated by reference into this Investor Pack. "
        "Executed packs should embed or annex the full PDFs."
    )
    if not atts:
        _p(
            doc,
            "No appendices selected in the wizard. Add vet / valuation / other under "
            f"_assets/horses/{evo}/documents/{{vet|valuation|…}}/.",
            italic=True,
        )
    else:
        for a in atts:
            doc.add_paragraph(
                f"[{a.get('category') or 'doc'}] {a.get('name')} — {a.get('path') or a.get('url') or ''}",
                style="List Bullet",
            )

    doc.add_paragraph("")
    _p(
        doc,
        "Counsel / founder: mark approved wording in this file, then freeze final PDF for investor send + e-sign.",
        italic=True,
        size=9,
    )
    _footer_line(doc, "evolutionstables.nz", f"{name} Syndicate · Investor Pack · DRAFT")
    _p(
        doc,
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} "
        f"from HLT Mission Control (investor pack template library)",
        size=8,
    )

    doc.save(str(out_path))
    return out_path


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(HERE), **kwargs)

    def log_message(self, fmt, *args):
        sys.stderr.write("[wizard] " + (fmt % args) + "\n")

    def _json(self, code: int, obj: dict):
        data = json.dumps(obj).encode("utf-8")
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(data)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        qs = parse_qs(parsed.query)

        if path == "/api/attachments":
            slug = (qs.get("slug") or [""])[0]
            if not slug:
                return self._json(400, {"error": "slug required"})
            return self._json(200, {"slug": slug, "attachments": list_attachments(slug)})

        if path.startswith("/horse-assets/"):
            rel = path[len("/horse-assets/") :]
            return self._serve_file(ASSETS_HORSES / rel)

        if path.startswith("/web-public/"):
            rel = path[len("/web-public/") :]
            return self._serve_file(WEB_PUBLIC / rel)

        return super().do_GET()

    def do_POST(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path)

        if path == "/api/upload":
            return self._handle_upload()
        if path == "/api/export-docx":
            return self._handle_docx()
        if path == "/api/approve-pack":
            return self._handle_approve()
        self.send_error(404, "Unknown API")

    def _handle_upload(self):
        ctype, pdict = cgi.parse_header(self.headers.get("Content-Type", ""))
        if ctype != "multipart/form-data":
            return self._json(400, {"error": "expected multipart/form-data"})
        pdict["boundary"] = bytes(pdict["boundary"], "utf-8")
        pdict["CONTENT-LENGTH"] = int(self.headers.get("Content-Length", 0))
        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": self.headers.get("Content-Type"),
            },
        )
        slug = form.getvalue("slug") or ""
        category = (form.getvalue("category") or "other").strip().lower()
        if category not in CATEGORIES:
            category = "other"
        if not slug:
            return self._json(400, {"error": "slug required"})
        if "file" not in form:
            return self._json(400, {"error": "file required"})
        fileitem = form["file"]
        if not getattr(fileitem, "file", None):
            return self._json(400, {"error": "empty file"})

        dest_dir = docs_root(slug) / category
        dest_dir.mkdir(parents=True, exist_ok=True)
        fname = safe_filename(fileitem.filename or "upload.pdf")
        # avoid overwrite
        dest = dest_dir / fname
        if dest.exists():
            stem, suf = dest.stem, dest.suffix
            dest = dest_dir / f"{stem}-{datetime.now().strftime('%H%M%S')}{suf}"
        data = fileitem.file.read()
        dest.write_bytes(data)

        rel = dest.relative_to(ASSETS_HORSES)
        url = "/horse-assets/" + "/".join(quote_seg(p) for p in rel.parts)
        return self._json(
            200,
            {
                "ok": True,
                "saved": str(dest.relative_to(ROOT)),
                "attachment": {
                    "category": category,
                    "name": dest.name,
                    "path": str(dest.relative_to(ROOT)),
                    "url": url,
                    "label": f"[{category}] {dest.name}",
                },
            },
        )

    def _handle_docx(self):
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length) if length else b"{}"
        try:
            payload = json.loads(raw.decode("utf-8"))
        except Exception as e:
            return self._json(400, {"error": f"invalid json: {e}"})
        try:
            path = build_docx(payload)
        except Exception as e:
            return self._json(500, {"error": str(e)})
        rel = path.relative_to(ASSETS_HORSES)
        url = "/horse-assets/" + "/".join(quote_seg(p) for p in rel.parts)
        return self._json(
            200,
            {
                "ok": True,
                "saved": str(path.relative_to(ROOT)),
                "url": url,
                "name": path.name,
            },
        )

    def _handle_approve(self):
        """
        Freeze pack as LIVE (good to go).
        Body JSON:
          { "slug": "...", "source_saved": "_assets/.../DRAFT.docx" optional,
            "payload": {...} optional — re-export first if no source,
            "note": "optional" }
        Multipart alternative: slug + file (counsel-edited DOCX) → becomes LIVE.
        """
        ctype = self.headers.get("Content-Type", "")
        note = ""
        slug = ""
        source: Path | None = None
        payload = None

        if ctype.startswith("multipart/form-data"):
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": ctype,
                },
            )
            slug = form.getvalue("slug") or ""
            note = form.getvalue("note") or ""
            if "file" in form and getattr(form["file"], "file", None):
                fileitem = form["file"]
                data = fileitem.file.read()
                if not slug:
                    return self._json(400, {"error": "slug required"})
                out_dir = docs_root(slug) / "investor-packs"
                out_dir.mkdir(parents=True, exist_ok=True)
                # save uploaded as staging then freeze
                staging = out_dir / safe_filename(fileitem.filename or "counsel-edit.docx")
                if not str(staging).lower().endswith(".docx"):
                    staging = out_dir / (staging.name + ".docx")
                staging.write_bytes(data)
                source = staging
        else:
            length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(length) if length else b"{}"
            try:
                body = json.loads(raw.decode("utf-8"))
            except Exception as e:
                return self._json(400, {"error": f"invalid json: {e}"})
            slug = body.get("slug") or ""
            note = body.get("note") or ""
            payload = body.get("payload")
            src = body.get("source_saved") or body.get("source") or ""
            if src:
                p = ROOT / src if not Path(src).is_absolute() else Path(src)
                if p.is_file():
                    source = p

        if not slug:
            # try from payload
            if payload and isinstance(payload, dict):
                slug = (payload.get("horse") or {}).get("slug") or ""
        if not slug:
            return self._json(400, {"error": "slug required"})

        if source is None:
            if payload:
                try:
                    source = build_docx(payload)
                except Exception as e:
                    return self._json(500, {"error": f"export failed: {e}"})
            else:
                return self._json(
                    400,
                    {
                        "error": "No draft to approve. Export DOCX first, or send payload/source_saved, or upload edited file.",
                    },
                )

        if not source.is_file():
            return self._json(400, {"error": f"source not found: {source}"})

        # path safety: must be under _assets/horses
        try:
            source_res = source.resolve()
            if not str(source_res).startswith(str(ASSETS_HORSES.resolve()) + os.sep):
                return self._json(403, {"error": "source must be under _assets/horses"})
        except Exception:
            return self._json(400, {"error": "bad source path"})

        out_dir = docs_root(slug) / "investor-packs"
        out_dir.mkdir(parents=True, exist_ok=True)
        live_name = f"Investor-Pack-{slug}-LIVE.docx"
        live_path = out_dir / live_name
        # archive previous LIVE if present
        if live_path.exists():
            stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
            archive = out_dir / f"Investor-Pack-{slug}-LIVE-archived-{stamp}.docx"
            live_path.replace(archive)

        import shutil

        shutil.copy2(source, live_path)

        meta = {
            "status": "LIVE",
            "slug": slug,
            "evo_slug": evo_slug(slug),
            "approved_at": datetime.now(timezone.utc).isoformat(),
            "live_docx": str(live_path.relative_to(ROOT)),
            "source_docx": str(source.relative_to(ROOT)) if str(source).startswith(str(ROOT)) else str(source),
            "note": note,
            "marketplace_note": (
                "Document pack frozen as LIVE. Marketplace soft/hard list is separate "
                "(campaign_status on Sheet/static). Set has_terms_sheet when PDFs are published for purchase."
            ),
        }
        meta_path = out_dir / f"Investor-Pack-{slug}-LIVE.json"
        meta_path.write_text(json.dumps(meta, indent=2) + "\n", encoding="utf-8")

        # also snapshot payload if provided
        if payload:
            (out_dir / f"Investor-Pack-{slug}-LIVE-payload.json").write_text(
                json.dumps(payload, indent=2) + "\n", encoding="utf-8"
            )

        rel = live_path.relative_to(ASSETS_HORSES)
        url = "/horse-assets/" + "/".join(quote_seg(p) for p in rel.parts)
        return self._json(
            200,
            {
                "ok": True,
                "status": "LIVE",
                "saved": str(live_path.relative_to(ROOT)),
                "meta": str(meta_path.relative_to(ROOT)),
                "url": url,
                "name": live_path.name,
                "message": "Pack approved and frozen as LIVE. Marketplace listing flags are still separate.",
            },
        )

    def _serve_file(self, file_path: Path):
        try:
            file_path = file_path.resolve()
        except Exception:
            self.send_error(400, "Bad path")
            return
        allowed = [ASSETS_HORSES.resolve(), WEB_PUBLIC.resolve()]
        if not any(str(file_path).startswith(str(a) + os.sep) or file_path == a for a in allowed):
            self.send_error(403, "Forbidden")
            return
        if not file_path.is_file():
            self.send_error(404, f"Not found: {file_path.name}")
            return
        ctype = mimetypes.guess_type(str(file_path))[0] or "application/octet-stream"
        data = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", ctype)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "private, max-age=30")
        self.end_headers()
        self.wfile.write(data)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--port", type=int, default=8765)
    ap.add_argument("--host", default="127.0.0.1")
    args = ap.parse_args()

    os.chdir(HERE)
    socketserver.TCPServer.allow_reuse_address = True
    with socketserver.TCPServer((args.host, args.port), Handler) as httpd:
        print(f"Investor Pack Wizard → http://{args.host}:{args.port}/")
        print(f"  uploads → _assets/horses/{{evo}}/documents/{{category}}/")
        print(f"  docx    → …/documents/investor-packs/")
        print("Ctrl+C to stop")
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("\nStopped.")


if __name__ == "__main__":
    main()
