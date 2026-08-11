"""Term Sheet DOCX generator.

Reads HLT + linked horse/owner/trainer/lease from SQLite,
produces a clean one-page term sheet in .docx format.
"""

from __future__ import annotations

import io
from datetime import date
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if TYPE_CHECKING:
    from sqlalchemy.orm import Session


def _set_cell(cell, text: str, bold: bool = False, size: int = 10):
    """Helper to style a table cell."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_term_sheet_docx(
    hlt_id: str,
    db: "Session",
) -> bytes:
    """Generate a term sheet DOCX for the given HLT. Returns raw bytes."""

    from admin.db import HLT as HLTORM, Lease as LeaseORM, Horse as HorseORM, Owner as OwnerORM, Trainer as TrainerORM

    hlt = db.query(HLTORM).filter_by(id=hlt_id).first()
    if not hlt:
        raise ValueError(f"HLT {hlt_id} not found")

    lease = db.query(LeaseORM).filter_by(lease_id=hlt.lease_id).first()
    horse = db.query(HorseORM).filter_by(microchip=hlt.horse_microchip).first()
    owner = db.query(OwnerORM).filter_by(id=hlt.owner_id).first()
    trainer = db.query(TrainerORM).filter_by(id=hlt.trainer_id).first()

    doc = Document()

    # ─── Styles ─────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ─── Header ─────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EVOLUTION STABLES — HORSE LEASE TERM SHEET")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x12, 0x12, 0x12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"HLT ID: {hlt_id}  |  Generated: {date.today().isoformat()}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ─── Horse Details ──────────────────────────────────────────────────────
    doc.add_heading("Horse Details", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("Name", horse.name if horse else "N/A"),
        ("Microchip", hlt.horse_microchip),
        ("Sire / Dam", f"{horse.sire_name or 'N/A'} / {horse.dam_name or 'N/A'}" if horse else "N/A"),
        ("Sex / Colour", f"{horse.sex or 'N/A'} / {horse.colour or 'N/A'}" if horse else "N/A"),
    ]
    for i, (label, value) in enumerate(rows):
        _set_cell(table.rows[i].cells[0], label, bold=True)
        _set_cell(table.rows[i].cells[1], value)

    doc.add_paragraph()

    # ─── Parties ────────────────────────────────────────────────────────────
    doc.add_heading("Parties", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    party_rows = [
        ("Owner", owner.name if owner else "N/A"),
        ("Trainer", trainer.name if trainer else "N/A"),
        ("Platform", "Evolution Stables Ltd"),
    ]
    for i, (label, value) in enumerate(party_rows):
        _set_cell(table.rows[i].cells[0], label, bold=True)
        _set_cell(table.rows[i].cells[1], value)

    doc.add_paragraph()

    # ─── Commercial Summary ─────────────────────────────────────────────────
    doc.add_heading("Commercial Summary", level=2)
    if lease:
        table = doc.add_table(rows=7, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        comm_rows = [
            ("Percent Leased", f"{lease.percent_leased}%"),
            ("Duration", f"{lease.duration_months} months"),
            ("Total Issuance Value", f"${lease.total_issuance_value_nzd:,.2f} NZD"),
            ("Token Count", str(lease.token_count)),
            ("Percent per Token", f"{lease.percent_per_token}%"),
            ("Token Price", f"${lease.token_price_nzd:,.2f} NZD"),
            ("Min Unit Size", f"{lease.min_unit_size}%"),
        ]
        for i, (label, value) in enumerate(comm_rows):
            _set_cell(table.rows[i].cells[0], label, bold=True)
            _set_cell(table.rows[i].cells[1], value)
    else:
        p = doc.add_paragraph("No lease attached.")
        p.italic = True

    doc.add_paragraph()

    # ─── Pricing Detail ─────────────────────────────────────────────────────
    doc.add_heading("Pricing Detail", level=2)
    if lease:
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        price_rows = [
            ("Price per 1% per month", f"${lease.price_per_1pct_per_month:,.2f}"),
            ("Price per 1% per year", f"${lease.price_per_1pct_per_year:,.2f}"),
            ("Monthly stake price", f"${lease.monthly_stake_price:,.2f}"),
            ("Annual stake price", f"${lease.annual_stake_price:,.2f}"),
        ]
        for i, (label, value) in enumerate(price_rows):
            _set_cell(table.rows[i].cells[0], label, bold=True)
            _set_cell(table.rows[i].cells[1], value)
    else:
        p = doc.add_paragraph("No lease attached.")
        p.italic = True

    doc.add_paragraph()

    # ─── Revenue Split ────────────────────────────────────────────────────────
    doc.add_heading("Revenue Split", level=2)
    if lease:
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        split_rows = [
            ("Investor Share", f"{lease.investor_share_percent}%"),
            ("Owner Share", f"{lease.owner_share_percent}%"),
            ("Platform Fee", f"{lease.platform_fee_percent}%"),
        ]
        for i, (label, value) in enumerate(split_rows):
            _set_cell(table.rows[i].cells[0], label, bold=True)
            _set_cell(table.rows[i].cells[1], value)
    else:
        p = doc.add_paragraph("No lease attached.")
        p.italic = True

    doc.add_paragraph()

    # ─── Legal Boilerplate ──────────────────────────────────────────────────
    doc.add_heading("Terms & Conditions", level=2)
    boiler = (
        "This term sheet is a summary of the proposed Horse Lease Token (HLT) offering. "
        "It is not a legally binding contract. A full Product Disclosure Statement (PDS) "
        "and Sale Agreement (SA) must be executed before any funds are accepted.\n\n"
        "All figures are estimates based on current pricing inputs and are subject to change. "
        "The governing body for this offering is the New Zealand Racing Board (NZTR).\n\n"
        "Termination: The lease may be terminated by mutual written agreement of all parties, "
        "or in accordance with the terms set out in the full Sale Agreement.\n\n"
        "Risk Warning: Horse racing involves significant financial risk. Past performance is not "
        "indicative of future results. Investors should seek independent legal and financial advice."
    )
    p = doc.add_paragraph(boiler)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # ─── Signature Blocks ───────────────────────────────────────────────────
    doc.add_heading("Signatures", level=2)
    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = "Table Grid"
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_headers = [("Party", "Signature", "Date")]
    sig_rows = [
        (owner.name if owner else "Owner", "____________________", "__________"),
        (trainer.name if trainer else "Trainer", "____________________", "__________"),
        ("Evolution Stables Ltd", "____________________", "__________"),
    ]
    for i, (party, sig, dt) in enumerate(sig_headers + sig_rows):
        _set_cell(sig_table.rows[i].cells[0], party, bold=(i == 0))
        _set_cell(sig_table.rows[i].cells[1], sig, bold=(i == 0))
        _set_cell(sig_table.rows[i].cells[2], dt, bold=(i == 0))

    # ─── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Evolution Stables Ltd  |  evolutionstables.nz  |  support@evolutionstables.nz")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ─── Output ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
