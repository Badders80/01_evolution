"""
Document Generation — Term Sheet, PDS, Syndicate Agreement

Generates DOCX files from HLT data, uploads to Cloud Storage,
and tracks them via DocumentRecord with section-by-section review.
"""

from flask import Request, jsonify
from google.cloud import firestore, storage
from datetime import datetime
import io
import re

from core.models import (
    DocumentRecord,
    DocumentRecordCreate,
    DocumentRecordUpdate,
    ReviewSection,
    build_default_sections,
    DOC_TYPE_SECTIONS,
)

_DB = None

def _get_db():
    global _DB
    if _DB is None:
        _DB = firestore.Client()
    return _DB

storage_client = storage.Client()
BUCKET_NAME = "evolution-horse-docs"

VALID_DOC_TYPES = ["term-sheet", "pds", "sa"]

# ─── Document Generation Handler ──────────────────────────────────────────────

def handle(request: Request, doc_type: str | None = None):
    """Generate a document from an HLT record and create a DocumentRecord."""
    if request.method != "POST":
        return jsonify({"error": "Method not allowed. Use POST."}), 405

    if doc_type not in VALID_DOC_TYPES:
        return jsonify({"error": f"Invalid doc type. Must be one of: {VALID_DOC_TYPES}"}), 400

    data = request.get_json(force=True)
    hlt_id = data.get("hlt_id")
    if not hlt_id:
        return jsonify({"error": "hlt_id is required"}), 400

    # Fetch HLT
    hlt_doc = _get_db().collection("hlts").document(hlt_id).get()
    if not hlt_doc.exists:
        return jsonify({"error": f"HLT {hlt_id} not found"}), 404

    hlt_data = hlt_doc.to_dict()

    # Fetch related entities
    horse_doc = _get_db().collection("horses").document(hlt_data["horse_microchip"]).get()
    horse_data = horse_doc.to_dict() if horse_doc.exists else {}

    owner_doc = _get_db().collection("owners").document(hlt_data["owner_id"]).get()
    owner_data = owner_doc.to_dict() if owner_doc.exists else {}

    trainer_doc = _get_db().collection("trainers").document(hlt_data["trainer_id"]).get()
    trainer_data = trainer_doc.to_dict() if trainer_doc.exists else {}

    # Generate document
    try:
        doc_bytes = generate_document(doc_type, hlt_data, horse_data, owner_data, trainer_data)
    except Exception as e:
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500

    # Upload to Cloud Storage
    bucket = storage_client.bucket(BUCKET_NAME)
    file_name = f"{hlt_id}/{doc_type}.docx"
    blob = bucket.blob(file_name)
    blob.upload_from_string(
        doc_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    gcs_url = f"gs://{BUCKET_NAME}/{file_name}"

    # Build canonical document ID
    lease_id = hlt_data.get("lease_id", "UNKNOWN")
    safe_lease = re.sub(r"[^A-Z0-9-]", "", lease_id.upper())
    safe_doc_type = doc_type.upper().replace("-", "")
    document_id = f"DOC-{safe_lease}-{safe_doc_type}"

    # Build DocumentRecord
    today = datetime.utcnow().date()
    sections = build_default_sections(doc_type)

    record = DocumentRecordCreate(
        document_id=document_id,
        lease_id=lease_id,
        horse_id=hlt_data["horse_microchip"],
        document_type=doc_type,
        document_version=1,
        document_date=today,
        source_reference="v1.0-template",
        file_path=gcs_url,
        is_current=True,
        notes=None,
        doc_review_status="draft",
        sections=sections,
    )

    # Write to Firestore
    doc_ref = _get_db().collection("documents").document(document_id)
    doc_ref.set(record.model_dump() | {
        "created_at": firestore.SERVER_TIMESTAMP,
        "updated_at": firestore.SERVER_TIMESTAMP,
    })

    # Update HLT document status (legacy field for backward compat)
    doc_field = doc_type.replace("-", "_")
    hlt_doc.reference.update({
        f"documents.{doc_field}.status": "pending",
        f"documents.{doc_field}.gcs_url": gcs_url,
        f"documents.{doc_field}.document_id": document_id,
        "updated_at": firestore.SERVER_TIMESTAMP,
    })

    return jsonify({
        "message": f"{doc_type} generated successfully",
        "gcs_url": gcs_url,
        "doc_type": doc_type,
        "hlt_id": hlt_id,
        "document_id": document_id,
        "sections": [s.section_name for s in sections],
    }), 200


# ─── Review Endpoint ──────────────────────────────────────────────────────────

def review(request: Request, document_id: str):
    """POST /docs/{document_id}/review — update section statuses and reviewer notes."""
    if request.method != "POST":
        return jsonify({"error": "Method not allowed. Use POST."}), 405

    data = request.get_json(force=True) or {}
    section_updates = data.get("sections", {})
    # e.g. {"horse_details": {"status": "approved", "reviewer_notes": "Looks good"}}

    doc_ref = _get_db().collection("documents").document(document_id)
    doc_snap = doc_ref.get()
    if not doc_snap.exists:
        return jsonify({"error": f"Document {document_id} not found"}), 404

    doc_data = doc_snap.to_dict()
    sections = doc_data.get("sections", [])

    updated_any = False
    for sec in sections:
        name = sec.get("section_name")
        if name in section_updates:
            upd = section_updates[name]
            if "status" in upd:
                sec["status"] = upd["status"]
                updated_any = True
            if "reviewer_notes" in upd:
                sec["reviewer_notes"] = upd["reviewer_notes"]
                updated_any = True

    # Auto-derive overall doc_review_status from sections
    all_approved = all(s.get("status") == "approved" for s in sections)
    any_rejected = any(s.get("status") == "rejected" for s in sections)
    any_needs_revision = any(s.get("status") == "needs_revision" for s in sections)

    new_status = doc_data.get("doc_review_status", "draft")
    if all_approved and len(sections) > 0:
        new_status = "approved"
    elif any_rejected:
        new_status = "rejected"
    elif any_needs_revision:
        new_status = "review"

    update_payload = {
        "sections": sections,
        "doc_review_status": new_status,
        "updated_at": firestore.SERVER_TIMESTAMP,
    }

    if updated_any:
        doc_ref.update(update_payload)

    return jsonify({
        "document_id": document_id,
        "doc_review_status": new_status,
        "sections": sections,
    }), 200


# ─── Get Document ─────────────────────────────────────────────────────────────

def get_document(request: Request, document_id: str):
    """GET /docs/{document_id} — fetch DocumentRecord with section statuses."""
    if request.method != "GET":
        return jsonify({"error": "Method not allowed. Use GET."}), 405

    doc_snap = _get_db().collection("documents").document(document_id).get()
    if not doc_snap.exists:
        return jsonify({"error": f"Document {document_id} not found"}), 404

    return jsonify(doc_snap.to_dict()), 200


# ─── List Documents by Lease ──────────────────────────────────────────────────

def list_by_lease(request: Request):
    """GET /docs?lease_id={id} — list all DocumentRecords for a lease."""
    if request.method != "GET":
        return jsonify({"error": "Method not allowed. Use GET."}), 405

    lease_id = request.args.get("lease_id") if hasattr(request, "args") else None
    if not lease_id:
        body = request.get_json(silent=True) or {}
        lease_id = body.get("lease_id")

    if not lease_id:
        return jsonify({"error": "lease_id is required (query param or body)."}), 400

    docs_query = _get_db().collection("documents").where("lease_id", "==", lease_id).stream()
    results = [d.to_dict() for d in docs_query]

    return jsonify({
        "lease_id": lease_id,
        "count": len(results),
        "documents": results,
    }), 200


# ─── DOCX Generation ──────────────────────────────────────────────────────────

def generate_document(doc_type: str, hlt_data: dict, horse_data: dict, owner_data: dict, trainer_data: dict) -> bytes:
    """Generate a DOCX file from HLT data."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    GOLD = RGBColor(0xD4, 0xA9, 0x64)
    BLACK = RGBColor(0x12, 0x12, 0x12)

    # Title
    title = doc.add_heading("", level=0)
    run = title.add_run(f"{doc_type.replace('-', ' ').title()}")
    run.font.size = Pt(24)
    run.font.color.rgb = BLACK

    # Subtitle
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Horse Lease Token — {horse_data.get('name', 'Unknown')}")
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD

    # Horse details
    doc.add_heading("Horse Details", level=1)
    details = [
        ("Name", horse_data.get("name", "—")),
        ("Microchip", horse_data.get("microchip", "—")),
        ("Life Number", horse_data.get("life_number", "—")),
        ("Foaling Date", str(horse_data.get("foaling_date", "—"))),
        ("Sex", horse_data.get("sex", "—")),
        ("Colour", horse_data.get("colour", "—")),
        ("Sire", horse_data.get("sire_name", "—")),
        ("Dam", horse_data.get("dam_name", "—")),
        ("Breeder", horse_data.get("breeder", "—")),
    ]
    table = doc.add_table(rows=len(details), cols=2)
    table.style = "Light List Accent 1"
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Owner details
    doc.add_heading("Owner Details", level=1)
    owner_details = [
        ("Name", owner_data.get("name", "—")),
        ("Type", owner_data.get("type", "—")),
        ("Email", owner_data.get("email", "—")),
    ]
    table = doc.add_table(rows=len(owner_details), cols=2)
    table.style = "Light List Accent 1"
    for i, (label, value) in enumerate(owner_details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Trainer details
    doc.add_heading("Trainer Details", level=1)
    trainer_details = [
        ("Name", trainer_data.get("name", "—")),
        ("Stable", trainer_data.get("stable_name", "—")),
        ("Location", trainer_data.get("location", "—")),
    ]
    table = doc.add_table(rows=len(trainer_details), cols=2)
    table.style = "Light List Accent 1"
    for i, (label, value) in enumerate(trainer_details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Lease terms
    doc.add_heading("Lease Terms", level=1)
    lease_details = [
        ("Lease Period", f"{hlt_data.get('lease_period_months', '—')} months"),
        ("Lease Start Date", str(hlt_data.get("lease_start_date", "—"))),
        ("Leasehold Stake", f"{hlt_data.get('leasehold_stake_percentage', '—')}%"),
        ("Investor Return", f"{hlt_data.get('investor_return_percentage', '—')}%"),
        ("Syndicate Price", f"${hlt_data.get('syndicate_price_cents', 0) / 100:,.2f} NZD"),
        ("Total Shares", str(hlt_data.get("shares_total", "—"))),
        ("Share Price", f"${hlt_data.get('share_price_cents', 0) / 100:,.2f} NZD"),
    ]
    table = doc.add_table(rows=len(lease_details), cols=2)
    table.style = "Light List Accent 1"
    for i, (label, value) in enumerate(lease_details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    run = footer.add_run("Evolution Stables — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = GOLD

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
